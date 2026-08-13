"""Word document extraction: mammoth for raw text, python-docx for structured tables.

Some annual reports (and their supporting notes) arrive as Word documents
rather than PDFs. This extractor keeps the rest of the pipeline — note
segmentation, line-item mapping — working identically regardless of source
format by returning the same ExtractedDocument/TableData shape as the PDF
extractor.
"""
import mammoth
from docx import Document as DocxDocument
from loguru import logger

from models.financial import ExtractedDocument, PageContent, TableData
from utils.text_utils import detect_company_name, detect_period


def _clean_table_rows(rows: list[list[str]]) -> list[list[str]]:
    """Remove blank rows and repeated header rows; strip whitespace from cells."""
    cleaned: list[list[str]] = []
    header_signature = None
    for i, row in enumerate(rows):
        stripped = [cell.strip().replace("\n", " ") if cell else "" for cell in row]
        if all(cell == "" for cell in stripped):
            continue
        if i == 0:
            header_signature = tuple(stripped)
        elif header_signature is not None and tuple(stripped) == header_signature:
            continue
        cleaned.append(stripped)
    return cleaned


def _extract_tables(docx_path: str) -> list[TableData]:
    tables: list[TableData] = []
    try:
        document = DocxDocument(docx_path)
    except Exception as e:
        logger.warning(f"python-docx could not open {docx_path} for table extraction: {e}")
        return tables

    for table in document.tables:
        try:
            raw_rows = [[cell.text for cell in row.cells] for row in table.rows]
        except Exception as e:
            logger.warning(f"Failed to read a table in {docx_path}: {e}")
            continue
        cleaned = _clean_table_rows(raw_rows)
        if not cleaned:
            continue
        headers = cleaned[0]
        rows = cleaned[1:]
        tables.append(TableData(headers=headers, rows=rows, page_num=1))
    return tables


def extract_docx(docx_path: str) -> ExtractedDocument:
    """Extract raw text and structured tables from a DOCX file.

    DOCX has no native page concept, so the whole document is returned as a
    single PageContent carrying every table found.
    """
    try:
        with open(docx_path, "rb") as f:
            result = mammoth.extract_raw_text(f)
        text = result.value
        for message in result.messages:
            logger.debug(f"mammoth: {message}")
    except Exception as e:
        logger.error(f"Failed to extract DOCX text {docx_path}: {e}")
        text = ""

    tables = _extract_tables(docx_path)

    return ExtractedDocument(
        pages=[PageContent(page_num=1, raw_text=text, tables=tables, ocr=False)],
        full_text=text,
        extraction_method="mammoth+python-docx",
        company_name=detect_company_name(text),
        period=detect_period(text),
        total_pages=1,
    )
