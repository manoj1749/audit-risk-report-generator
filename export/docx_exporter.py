"""Export the final audit report as a formatted DOCX working paper."""
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement

from export.charts import (
    generate_movements_chart,
    generate_ratios_chart,
    generate_risk_distribution_chart,
)
from models.report import AuditReport
from utils.timezone import format_ist

_RISK_FILL = {"High": "F8D7DA", "Medium": "FFF3CD", "Low": "D4EDDA"}
_DISCLAIMER_FILL = "F0F0F0"


def _shade_cell(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = "audit-risk-report-generator | Preliminary Review Only"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_chart(document: Document, png_bytes: bytes | None, width_inches: float = 6.0) -> None:
    if not png_bytes:
        return
    document.add_picture(io.BytesIO(png_bytes), width=Inches(width_inches))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_report_docx(report: AuditReport) -> bytes:
    document = Document()
    _add_footer(document)

    # Company name, prominent
    if report.company_name:
        company = document.add_paragraph()
        company_run = company.add_run(report.company_name.upper())
        company_run.bold = True
        company_run.font.size = Pt(14)

    # Title
    title = document.add_paragraph()
    title_run = title.add_run("RISK ASSESSMENT & DRAFT AUDIT OBSERVATIONS")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Subtitle
    if report.period:
        subtitle = document.add_paragraph()
        subtitle_run = subtitle.add_run(
            f"Financial statements for the period ended {report.period} — "
            f"cross-referenced to the notes to accounts"
        )
        subtitle_run.font.size = Pt(11)
        subtitle_run.italic = True

    generated = document.add_paragraph()
    generated.add_run(f"Generated: {format_ist(report.generated_at)}").italic = True

    document.add_paragraph()

    # Section 1: Purpose & Basis
    document.add_heading("1. Purpose & Basis", level=2)
    document.add_paragraph(
        f"{report.disclaimer} Figures are in ₹ lakh unless stated otherwise."
    )
    disclaimer_table = document.add_table(rows=1, cols=1)
    disclaimer_cell = disclaimer_table.rows[0].cells[0]
    disclaimer_cell.text = report.disclaimer
    _shade_cell(disclaimer_cell, _DISCLAIMER_FILL)

    document.add_paragraph()

    # Section 2: Summary of Risk Ratings
    document.add_heading("2. Summary of Risk Ratings", level=2)
    summary_table = document.add_table(rows=1, cols=3)
    summary_table.style = "Light Grid Accent 1"
    hdr = summary_table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Risk Rating", "Count", "Focus Areas"
    for rating in ["High", "Medium", "Low"]:
        row = summary_table.add_row().cells
        row[0].text = rating
        row[1].text = str(report.summary.get(rating, 0))
        areas = sorted({o.area for o in report.observations if o.risk_rating == rating})
        row[2].text = ", ".join(areas) if areas else "-"
        for cell in row:
            _shade_cell(cell, _RISK_FILL[rating])

    document.add_paragraph()
    _add_chart(document, generate_risk_distribution_chart(report.summary), width_inches=4.5)

    document.add_paragraph()

    # Section 3: Detailed Risk Areas & Draft Audit Observations
    document.add_heading("3. Detailed Risk Areas & Draft Audit Observations", level=2)
    obs_table = document.add_table(rows=1, cols=7)
    obs_table.style = "Light Grid Accent 1"
    headers = ["#", "Area / Process", "Audit Observation", "Note Ref", "Risk Rating",
               "Standard Reference", "Recommendation"]
    for i, h in enumerate(headers):
        obs_table.rows[0].cells[i].text = h

    for i, obs in enumerate(report.observations, 1):
        row = obs_table.add_row().cells
        row[0].text = str(i)
        row[1].text = obs.area
        row[2].text = obs.observation
        row[3].text = obs.note_ref
        row[4].text = obs.risk_rating
        row[5].text = obs.standard_reference
        row[6].text = obs.recommendation
        for cell in row:
            _shade_cell(cell, _RISK_FILL[obs.risk_rating])

    document.add_paragraph()

    # Section 4: Key Financial Movements
    document.add_heading("4. Key Financial Movements", level=2)
    mov_table = document.add_table(rows=1, cols=4)
    mov_table.style = "Light Grid Accent 1"
    mov_headers = ["Line Item", "Current", "Prior", "Change %"]
    for i, h in enumerate(mov_headers):
        mov_table.rows[0].cells[i].text = h
    for m in report.key_movements:
        row = mov_table.add_row().cells
        row[0].text = m.display_label
        row[1].text = f"{m.current:,.0f}" if m.current is not None else "-"
        row[2].text = f"{m.prior:,.0f}" if m.prior is not None else "-"
        row[3].text = f"{m.pct_change:+.1f}%" if m.pct_change is not None else "-"

    document.add_paragraph()
    _add_chart(document, generate_movements_chart(report.key_movements))

    document.add_paragraph()

    # Section 5: Computed Ratios
    document.add_heading("5. Computed Ratios", level=2)
    ratios = report.ratios
    ratio_rows = [
        ("Current Ratio", ratios.current_ratio_current, ratios.current_ratio_prior),
        ("Quick Ratio", ratios.quick_ratio_current, ratios.quick_ratio_prior),
        ("Debt-Equity Ratio", ratios.debt_equity_current, ratios.debt_equity_prior),
        ("Interest Coverage", ratios.interest_coverage_current, ratios.interest_coverage_prior),
        ("Debtor Days", ratios.debtor_days_current, ratios.debtor_days_prior),
        ("Creditor Days", ratios.creditor_days_current, ratios.creditor_days_prior),
        ("Net Profit Margin %", ratios.net_profit_margin_current, ratios.net_profit_margin_prior),
        ("ROCE %", ratios.roce_current, ratios.roce_prior),
        ("CFO / PAT", ratios.cfo_to_pat_current, ratios.cfo_to_pat_prior),
    ]
    ratio_table = document.add_table(rows=1, cols=3)
    ratio_table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Ratio", "Current", "Prior"]):
        ratio_table.rows[0].cells[i].text = h
    for name, curr, prior in ratio_rows:
        row = ratio_table.add_row().cells
        row[0].text = name
        row[1].text = f"{curr:.2f}" if curr is not None else "-"
        row[2].text = f"{prior:.2f}" if prior is not None else "-"

    document.add_paragraph()
    _add_chart(document, generate_ratios_chart(ratios))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
